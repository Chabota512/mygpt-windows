"""
My_GPT 4 Students — Local Python Backend
=========================================

A self-contained FastAPI server that the React frontend talks to on localhost.
Runs fully OFFLINE on a Windows PC. Stores data in a local SQLite file and
keeps uploaded files / generated documents on disk next to this script.

Quick start (Windows):

    python -m venv .venv
    .venv\\Scripts\\activate
    pip install -r requirements.txt
    python main.py

Then run the React frontend. It points at http://localhost:8000 by default.

Plugging in Ollama (real offline LLM) — see README.md.
"""

from __future__ import annotations

import os
import sys
import re
import sqlite3
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List, Optional

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import uvicorn


# ──────────────────────────────────────────────────────────────────────
# Paths & DB
# ──────────────────────────────────────────────────────────────────────
# "Head on PC, body wherever you want" model:
#   - Code lives next to this file (small).
#   - Data (uploads, generated docs, chat database, profile) lives in DATA_DIR.
#
# Where DATA_DIR comes from, in order of priority:
#   1. Environment variable  MYGPT_DATA_DIR=E:\MyGPTData
#   2. data_location.txt    sitting next to this script (one line: a folder path)
#   3. Default:             ./data  (next to this script)
#
# Users with low disk space can simply edit data_location.txt and point it at
# an external drive — no code changes needed.
# ──────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent
LOCATION_FILE = ROOT / "data_location.txt"
MODEL_LOCATION_FILE = ROOT / "model_location.txt"


def resolve_data_dir() -> Path:
    env = os.environ.get("MYGPT_DATA_DIR", "").strip()
    if env:
        return Path(env).expanduser().resolve()
    if LOCATION_FILE.exists():
        line = LOCATION_FILE.read_text("utf-8").strip().splitlines()
        for raw in line:
            raw = raw.strip()
            if raw and not raw.startswith("#"):
                return Path(raw).expanduser().resolve()
    return (ROOT / "data").resolve()


def resolve_model_dir() -> Path:
    env = os.environ.get("MYGPT_MODEL_DIR", "").strip()
    if env:
        return Path(env).expanduser().resolve()
    if MODEL_LOCATION_FILE.exists():
        line = MODEL_LOCATION_FILE.read_text("utf-8").strip().splitlines()
        for raw in line:
            raw = raw.strip()
            if raw and not raw.startswith("#"):
                return Path(raw).expanduser().resolve()
    return (ROOT / "models").resolve()


DATA_DIR = resolve_data_dir()
MODEL_DIR = resolve_model_dir()
try:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    MODEL_DIR.mkdir(parents=True, exist_ok=True)
except OSError as e:
    raise SystemExit(
        f"Could not create data folder at {DATA_DIR}.\n"
        f"  Reason: {e}\n"
        f"  Fix: edit {LOCATION_FILE} and set a folder you can write to "
        f"(for example E:\\MyGPTData on an external drive)."
    )

UPLOAD_DIR = DATA_DIR / "uploads"
EXPORT_DIR = DATA_DIR / "exports"
DB_PATH    = DATA_DIR / "studentai.db"
UPLOAD_DIR.mkdir(exist_ok=True)
EXPORT_DIR.mkdir(exist_ok=True)

print(f"[My_GPT] Data folder: {DATA_DIR}")
print(f"[My_GPT] Model folder: {MODEL_DIR}")


def _dir_size(p: Path) -> int:
    total = 0
    try:
        for entry in p.rglob("*"):
            try:
                if entry.is_file():
                    total += entry.stat().st_size
            except OSError:
                pass
    except OSError:
        pass
    return total


def _human_bytes(n: int) -> str:
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if n < 1024 or unit == "TB":
            return f"{n:.1f} {unit}" if unit != "B" else f"{n} B"
        n /= 1024
    return f"{n:.1f} TB"


def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db() -> None:
    with db() as c:
        c.executescript(
            """
            CREATE TABLE IF NOT EXISTS sessions (
                id          TEXT PRIMARY KEY,
                label       TEXT NOT NULL,
                created_at  TEXT NOT NULL,
                updated_at  TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS messages (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id  TEXT NOT NULL REFERENCES sessions(id) ON DELETE CASCADE,
                role        TEXT NOT NULL,
                content     TEXT NOT NULL,
                created_at  TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS memory_items (
                id          TEXT PRIMARY KEY,
                filename    TEXT NOT NULL,
                kind        TEXT NOT NULL,             -- 'pdf' | 'image' | 'doc'
                size_bytes  INTEGER NOT NULL,
                path        TEXT NOT NULL,
                text        TEXT NOT NULL DEFAULT '',  -- extracted text for search
                created_at  TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS documents (
                id          TEXT PRIMARY KEY,
                title       TEXT NOT NULL,
                format      TEXT NOT NULL,             -- 'docx' | 'pdf'
                path        TEXT NOT NULL,
                size_bytes  INTEGER NOT NULL,
                prompt      TEXT NOT NULL DEFAULT '',
                created_at  TEXT NOT NULL
            );
            """
        )

        # Seed one welcome session if DB is fresh
        cur = c.execute("SELECT COUNT(*) AS n FROM sessions")
        if cur.fetchone()["n"] == 0:
            sid = uuid.uuid4().hex[:10]
            now = datetime.utcnow().isoformat()
            c.execute(
                "INSERT INTO sessions (id, label, created_at, updated_at) VALUES (?,?,?,?)",
                (sid, "Welcome", now, now),
            )


init_db()


# ──────────────────────────────────────────────────────────────────────
# App + CORS
# ──────────────────────────────────────────────────────────────────────
app = FastAPI(title="My_GPT 4 Students — Local Backend", version="0.2.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ──────────────────────────────────────────────────────────────────────
# Schemas
# ──────────────────────────────────────────────────────────────────────
class ChatMsg(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    attachment_ids: Optional[List[str]] = None  # memory_item ids (images) to feed the vision model


class ChatResponse(BaseModel):
    reply: str
    timestamp: str
    model: str
    session_id: str
    stages: List[str] = []          # human-readable steps the router took
    model_used: Optional[str] = None  # final model that produced the reply


class SessionOut(BaseModel):
    id: str
    label: str
    time: str  # "Now", "Today", date — friendly relative label


class SessionCreate(BaseModel):
    label: Optional[str] = "New Chat"


class SessionRename(BaseModel):
    label: str


class MemoryItemOut(BaseModel):
    id: str
    filename: str
    kind: str
    size_bytes: int
    created_at: str


class SearchRequest(BaseModel):
    query: str


class SearchHit(BaseModel):
    id: str
    type: str        # 'pdf' | 'image' | 'doc'
    title: str
    path: str
    snippet: str
    meta: str
    relevance: int


class SearchResponse(BaseModel):
    query: str
    results: List[SearchHit]


class GenerateDocRequest(BaseModel):
    prompt: str
    title: Optional[str] = None
    format: str = "docx"     # 'docx' | 'pdf'
    session_id: Optional[str] = None


class DocumentOut(BaseModel):
    id: str
    title: str
    format: str
    size_bytes: int
    created_at: str
    download_url: str
    sections: List[str]


# ──────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────
def _kind_for(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}:
        return "image"
    return "doc"


def _friendly_time(iso_ts: str) -> str:
    try:
        ts = datetime.fromisoformat(iso_ts.rstrip("Z"))
    except Exception:
        return ""
    delta = datetime.utcnow() - ts
    if delta.total_seconds() < 60:
        return "Now"
    if delta.total_seconds() < 3600:
        return f"{int(delta.total_seconds() // 60)}m ago"
    if delta.days == 0:
        return "Today"
    if delta.days == 1:
        return "Yesterday"
    if delta.days < 7:
        return ts.strftime("%a")
    return ts.strftime("%b %d")


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # lazy import
        reader = PdfReader(str(path))
        chunks = []
        for page in reader.pages[:50]:  # cap for safety
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(chunks).strip()
    except Exception as e:
        print(f"[pdf] extract failed for {path.name}: {e}")
        return ""


def _bump_session(session_id: str) -> None:
    with db() as c:
        c.execute(
            "UPDATE sessions SET updated_at = ? WHERE id = ?",
            (datetime.utcnow().isoformat(), session_id),
        )


# ──────────────────────────────────────────────────────────────────────
# Specialist Team — Task Router
# ──────────────────────────────────────────────────────────────────────
# We run THREE small, specialised local models through Ollama. Only ONE
# is in RAM at a time (OLLAMA_MAX_LOADED_MODELS=1) so an 8 GB machine
# can swap between them safely.
#
#   Vision     → describes any uploaded image in technical detail
#   Reasoning  → math / physics / code (Mechatronics-friendly)
#   Writer     → formats the final answer cleanly
#
# Defaults match the user's spec but every model + the Ollama host can
# be overridden with environment variables, so you can swap to whatever
# you've actually pulled with `ollama pull ...` without touching code.
# ──────────────────────────────────────────────────────────────────────
os.environ.setdefault("OLLAMA_MAX_LOADED_MODELS", "1")

OLLAMA_HOST           = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_VISION_MODEL   = os.environ.get("OLLAMA_VISION_MODEL",   "qwen3.5:0.8b")
OLLAMA_REASONING_MODEL = os.environ.get("OLLAMA_REASONING_MODEL", "phi4-mini")
OLLAMA_WRITER_MODEL   = os.environ.get("OLLAMA_WRITER_MODEL",   "llama3.2:1b")
OLLAMA_NUM_CTX        = int(os.environ.get("OLLAMA_NUM_CTX", "4096"))
# How many critique → revise cycles the reasoning specialist runs on its
# own draft before handing it to the writer. 0 disables debate (single shot).
# 1 is the sweet spot for an 8 GB laptop (one extra critique + one revision,
# all done while the reasoning model is still in RAM — no extra swap cost).
OLLAMA_DEBATE_ROUNDS  = max(0, int(os.environ.get("OLLAMA_DEBATE_ROUNDS", "1")))

_REASONING_HINTS = re.compile(
    r"(?ix)"
    r"(?:^|\b)("
    r"calculate|solve|prove|integrate|derivative|differentiat|equation|equations?|"
    r"matrix|matrices|laplace|fourier|transform|pid|tuning|circuit|voltage|current|"
    r"resistance|impedance|torque|force|velocity|acceleration|kinematic|dynamic|"
    r"physics|chemistry|formula|theorem|algorithm|big[- ]?o|complexity|"
    r"function|variable|loop|class|method|stack|trace|debug|compile|runtime|"
    r"python|java(?:script)?|typescript|c\+\+|c\#|rust|sql|regex"
    r")\b"
)
_CODE_FENCE = re.compile(r"```|`[^`]+`")
_MATH_SYMBOLS = re.compile(r"[=+\-*/^∫∑√π°][\s\d(]|[0-9]\s*[+\-*/^=]\s*[0-9]")


def _is_reasoning_task(message: str) -> bool:
    if _CODE_FENCE.search(message):
        return True
    if _MATH_SYMBOLS.search(message):
        return True
    return bool(_REASONING_HINTS.search(message))


def _ollama_chat(
    model: str,
    messages: list[dict],
    images: Optional[list[str]] = None,
) -> str:
    """Single Ollama call with our standard options. Raises on failure."""
    import ollama  # lazy import so dev/test can run without it
    client = ollama.Client(host=OLLAMA_HOST)
    if images and messages:
        # Attach images to the LAST user message (Ollama convention)
        for m in reversed(messages):
            if m.get("role") == "user":
                m["images"] = images
                break
    out = client.chat(
        model=model,
        messages=messages,
        options={"num_ctx": OLLAMA_NUM_CTX},
        stream=False,
    )
    return (out.get("message") or {}).get("content", "").strip()


def _run_reasoning_debate(
    base_history: list[dict],
    user_block: str,
    rounds: int,
    stages: list[str],
) -> str:
    """
    Have the reasoning specialist talk to itself: draft → critique → revise.

    Why same model for all calls?
      - We're locked to OLLAMA_MAX_LOADED_MODELS=1 on an 8 GB box.
      - Keeping the SAME model for every debate turn means it's only loaded
        ONCE. No swap cost between rounds.
      - The model plays different roles via different system prompts —
        a clean, well-known agent-debate pattern.

    Returns the final, revised reasoning answer.
    """
    sys_solver = (
        "You are a careful STEM reasoning assistant for a Mechatronics "
        "student. Show working step-by-step. Keep units. Prefer correctness "
        "over brevity."
    )
    sys_critic = (
        "You are a strict reviewer. Read the draft answer below and find "
        "concrete problems: math errors, wrong units, missing steps, wrong "
        "assumptions, unclear bits, or anything that could mislead a "
        "student. List ONLY the issues, numbered, one per line. If the "
        "draft is genuinely correct and complete, reply with the single "
        "word: NONE."
    )
    sys_reviser = (
        "You are the same reasoning assistant. Rewrite your previous answer "
        "to fix EVERY issue listed by the reviewer. Keep what was correct. "
        "Do not apologise or mention the review — just produce the better "
        "answer."
    )

    # ── Round 0 · initial draft ────────────────────────────────
    draft = _ollama_chat(
        OLLAMA_REASONING_MODEL,
        base_history + [
            {"role": "system", "content": sys_solver},
            {"role": "user", "content": user_block},
        ],
    )

    # ── Rounds 1..N · critique + revise ────────────────────────
    for i in range(1, rounds + 1):
        stages.append(f"Critiquing draft (round {i})…")
        critique = _ollama_chat(
            OLLAMA_REASONING_MODEL,
            [
                {"role": "system", "content": sys_critic},
                {
                    "role": "user",
                    "content": (
                        f"Original question:\n{user_block}\n\n"
                        f"Draft answer:\n{draft}"
                    ),
                },
            ],
        )

        # Critic says the draft is fine — stop early, don't waste a swap.
        if critique.strip().upper().startswith("NONE"):
            stages.append("Reviewer approved the draft — stopping debate.")
            break

        stages.append(f"Revising answer (round {i})…")
        draft = _ollama_chat(
            OLLAMA_REASONING_MODEL,
            base_history + [
                {"role": "system", "content": sys_reviser},
                {"role": "user", "content": user_block},
                {"role": "assistant", "content": draft},
                {
                    "role": "user",
                    "content": (
                        "Reviewer found these issues — fix them all:\n"
                        f"{critique}"
                    ),
                },
            ],
        )

    return draft


def _resolve_image_paths(attachment_ids: list[str]) -> list[str]:
    if not attachment_ids:
        return []
    paths: list[str] = []
    with db() as c:
        q_marks = ",".join("?" * len(attachment_ids))
        rows = c.execute(
            f"SELECT path, kind FROM memory_items WHERE id IN ({q_marks})",
            attachment_ids,
        ).fetchall()
    for r in rows:
        if r["kind"] == "image" and Path(r["path"]).exists():
            paths.append(r["path"])
    return paths


def _route_task(
    message: str,
    history: List[ChatMsg],
    image_paths: Optional[list[str]] = None,
) -> tuple[str, list[str], str]:
    """
    Picks the right specialist(s) for this task and returns
    (final_reply, stages, final_model_used).

    Falls back to the offline stub if Ollama isn't installed / reachable.
    """
    image_paths = image_paths or []
    stages: list[str] = []

    try:
        import ollama  # noqa: F401
    except ImportError:
        return (
            f"(offline stub) I received: \"{message.strip()}\".\n\n"
            "Install Ollama and `pip install ollama` to enable the local AI "
            "specialists — see README.md.",
            ["Offline stub (no Ollama installed)"],
            "stub",
        )

    base_history = [{"role": m.role, "content": m.content} for m in history]

    try:
        # ── Step 1 · Vision (only if an image is attached) ─────────
        vision_desc = ""
        if image_paths:
            stages.append(f"Analyzing image with {OLLAMA_VISION_MODEL}…")
            vision_desc = _ollama_chat(
                OLLAMA_VISION_MODEL,
                [
                    {
                        "role": "system",
                        "content": (
                            "You are a technical vision assistant. Describe the "
                            "image in precise, factual detail: text shown, "
                            "diagrams, equations, components, measurements. Be "
                            "concise but complete. Do not speculate."
                        ),
                    },
                    {"role": "user", "content": message or "Describe this image."},
                ],
                images=image_paths,
            )

        # ── Step 2 · Reasoning, with optional self-debate ──────────
        use_reasoning = _is_reasoning_task(message) or bool(vision_desc)
        reasoning_out = ""
        if use_reasoning:
            stages.append(f"Drafting answer with {OLLAMA_REASONING_MODEL}…")
            user_block = message
            if vision_desc:
                user_block = (
                    f"Image description (from vision specialist):\n{vision_desc}\n\n"
                    f"User question:\n{message}"
                )
            reasoning_out = _run_reasoning_debate(
                base_history, user_block, OLLAMA_DEBATE_ROUNDS, stages,
            )

        # ── Step 3 · Writer / formatter ─────────────────────────────
        # If we already have a reasoning answer, polish it. The writer also
        # acts as a final "clarity reviewer" — it can flag confusion by
        # prefixing its reply with `[NEEDS_REWORK]`, which sends us back to
        # the reasoning specialist for one more pass. This is the cross-model
        # conversation: writer ↔ reasoning, capped at 1 bounce so we don't
        # thrash the RAM.
        stages.append(f"Formatting with {OLLAMA_WRITER_MODEL}…")
        if reasoning_out:
            writer_system = (
                "You are an academic writing assistant. Take the draft below "
                "and rewrite it as a clean, well-structured answer for a "
                "student. Preserve every number, equation and step. Use short "
                "paragraphs, headings or bullets where helpful. Do NOT add "
                "new facts.\n\n"
                "If — and ONLY if — the draft is so unclear or contradictory "
                "that you cannot polish it without inventing content, reply "
                "with EXACTLY this format and nothing else:\n"
                "[NEEDS_REWORK]\n<one short sentence saying what is unclear>"
            )
            writer_msgs = base_history + [
                {"role": "system", "content": writer_system},
                {"role": "user", "content": f"Draft to polish:\n\n{reasoning_out}"},
            ]
        else:
            writer_msgs = base_history + [
                {
                    "role": "system",
                    "content": (
                        "You are a friendly study assistant for a student. "
                        "Answer clearly and concisely in plain language."
                    ),
                },
                {"role": "user", "content": message},
            ]

        final = _ollama_chat(OLLAMA_WRITER_MODEL, writer_msgs)

        # ── Step 4 · Optional bounce-back if writer asked for clarity ──
        if (
            reasoning_out
            and final.lstrip().startswith("[NEEDS_REWORK]")
        ):
            feedback = final.split("\n", 1)[1].strip() if "\n" in final else "Make it clearer and more complete."
            stages.append("Writer asked for clarity — sending back to reasoning…")
            reasoning_out = _ollama_chat(
                OLLAMA_REASONING_MODEL,
                base_history + [
                    {
                        "role": "system",
                        "content": (
                            "You are the reasoning assistant. Your previous "
                            "answer was unclear to the writing assistant. "
                            "Rewrite it more clearly while preserving every "
                            "correct step, number and unit."
                        ),
                    },
                    {"role": "user", "content": message},
                    {"role": "assistant", "content": reasoning_out},
                    {
                        "role": "user",
                        "content": f"Writer's feedback: {feedback}\n\nPlease rewrite.",
                    },
                ],
            )
            stages.append(f"Re-formatting with {OLLAMA_WRITER_MODEL}…")
            final = _ollama_chat(
                OLLAMA_WRITER_MODEL,
                base_history + [
                    {
                        "role": "system",
                        "content": (
                            "You are an academic writing assistant. Polish "
                            "this revised draft into a clean answer for a "
                            "student. Preserve every number, equation and "
                            "step. Do not add new facts."
                        ),
                    },
                    {"role": "user", "content": f"Draft to polish:\n\n{reasoning_out}"},
                ],
            )

        if not final:
            # Writer produced nothing — fall back to the upstream output
            final = reasoning_out or vision_desc or "(no response)"
        return final, stages, OLLAMA_WRITER_MODEL

    except Exception as e:
        print(f"[router] Ollama call failed: {e}")
        # Graceful fallback so the UI never crashes
        return (
            "I couldn't reach the local AI just now.\n\n"
            f"• Make sure Ollama is running (`ollama serve`).\n"
            f"• Make sure you've pulled: `{OLLAMA_VISION_MODEL}`, "
            f"`{OLLAMA_REASONING_MODEL}`, `{OLLAMA_WRITER_MODEL}`.\n"
            f"• Technical detail: {e}",
            stages + ["Failed — see message"],
            "error",
        )


# ──────────────────────────────────────────────────────────────────────
# Root / health
# ──────────────────────────────────────────────────────────────────────
@app.get("/")
def root():
    return {"app": "My_GPT 4 Students backend", "status": "ok"}


@app.get("/health")
def health():
    return {"status": "ok", "time": datetime.utcnow().isoformat() + "Z"}


# ──────────────────────────────────────────────────────────────────────
# Sessions
# ──────────────────────────────────────────────────────────────────────
@app.get("/sessions", response_model=List[SessionOut])
def list_sessions():
    with db() as c:
        rows = c.execute(
            "SELECT id, label, updated_at FROM sessions ORDER BY updated_at DESC"
        ).fetchall()
    return [
        SessionOut(id=r["id"], label=r["label"], time=_friendly_time(r["updated_at"]))
        for r in rows
    ]


@app.post("/sessions", response_model=SessionOut)
def create_session(body: SessionCreate):
    sid = uuid.uuid4().hex[:10]
    now = datetime.utcnow().isoformat()
    label = (body.label or "New Chat").strip() or "New Chat"
    with db() as c:
        c.execute(
            "INSERT INTO sessions (id, label, created_at, updated_at) VALUES (?,?,?,?)",
            (sid, label, now, now),
        )
    return SessionOut(id=sid, label=label, time="Now")


@app.patch("/sessions/{session_id}", response_model=SessionOut)
def rename_session(session_id: str, body: SessionRename):
    label = body.label.strip()
    if not label:
        raise HTTPException(400, "Label cannot be empty.")
    with db() as c:
        c.execute("UPDATE sessions SET label = ? WHERE id = ?", (label, session_id))
        row = c.execute(
            "SELECT id, label, updated_at FROM sessions WHERE id = ?", (session_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "Session not found.")
    return SessionOut(id=row["id"], label=row["label"], time=_friendly_time(row["updated_at"]))


@app.delete("/sessions/{session_id}")
def delete_session(session_id: str):
    with db() as c:
        c.execute("DELETE FROM sessions WHERE id = ?", (session_id,))
    return {"ok": True}


@app.get("/sessions/{session_id}/messages", response_model=List[ChatMsg])
def session_messages(session_id: str):
    with db() as c:
        rows = c.execute(
            "SELECT role, content FROM messages WHERE session_id = ? ORDER BY id ASC",
            (session_id,),
        ).fetchall()
    return [ChatMsg(role=r["role"], content=r["content"]) for r in rows]


# ──────────────────────────────────────────────────────────────────────
# Chat
# ──────────────────────────────────────────────────────────────────────
@app.post("/chat", response_model=ChatResponse)
def chat(req: ChatRequest):
    msg = (req.message or "").strip()
    if not msg:
        raise HTTPException(400, "Message cannot be empty.")

    # Ensure session exists (create on the fly if needed)
    sid = req.session_id
    with db() as c:
        if not sid or not c.execute(
            "SELECT 1 FROM sessions WHERE id = ?", (sid,)
        ).fetchone():
            sid = uuid.uuid4().hex[:10]
            now = datetime.utcnow().isoformat()
            label = (msg[:40] + ("…" if len(msg) > 40 else "")) or "New Chat"
            c.execute(
                "INSERT INTO sessions (id, label, created_at, updated_at) VALUES (?,?,?,?)",
                (sid, label, now, now),
            )

        # Pull last 20 messages as history for the LLM
        rows = c.execute(
            "SELECT role, content FROM messages WHERE session_id = ? "
            "ORDER BY id DESC LIMIT 20",
            (sid,),
        ).fetchall()
    history = [ChatMsg(role=r["role"], content=r["content"]) for r in reversed(rows)]

    image_paths = _resolve_image_paths(req.attachment_ids or [])
    reply, stages, model_used = _route_task(msg, history, image_paths)

    now_iso = datetime.utcnow().isoformat()
    with db() as c:
        c.execute(
            "INSERT INTO messages (session_id, role, content, created_at) VALUES (?,?,?,?)",
            (sid, "user", msg, now_iso),
        )
        c.execute(
            "INSERT INTO messages (session_id, role, content, created_at) VALUES (?,?,?,?)",
            (sid, "assistant", reply, now_iso),
        )
    _bump_session(sid)

    return ChatResponse(
        reply=reply,
        timestamp=datetime.now().strftime("%H:%M"),
        model=model_used,
        session_id=sid,
        stages=stages,
        model_used=model_used,
    )


# ──────────────────────────────────────────────────────────────────────
# Memory (uploads)
# ──────────────────────────────────────────────────────────────────────
@app.get("/memory", response_model=List[MemoryItemOut])
def list_memory():
    with db() as c:
        rows = c.execute(
            "SELECT id, filename, kind, size_bytes, created_at "
            "FROM memory_items ORDER BY created_at DESC"
        ).fetchall()
    return [MemoryItemOut(**dict(r)) for r in rows]


@app.post("/upload-pdf", response_model=MemoryItemOut)
async def upload_pdf(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(400, "No filename provided.")

    file_id = uuid.uuid4().hex[:12]
    safe = Path(file.filename).name
    target = UPLOAD_DIR / f"{file_id}__{safe}"
    contents = await file.read()
    target.write_bytes(contents)

    kind = _kind_for(safe)
    text = _extract_pdf_text(target) if kind == "pdf" else ""
    now = datetime.utcnow().isoformat()
    with db() as c:
        c.execute(
            "INSERT INTO memory_items (id, filename, kind, size_bytes, path, text, created_at) "
            "VALUES (?,?,?,?,?,?,?)",
            (file_id, safe, kind, len(contents), str(target), text, now),
        )

    print(f"[upload] {safe}  ({len(contents)} bytes, kind={kind})")
    return MemoryItemOut(
        id=file_id,
        filename=safe,
        kind=kind,
        size_bytes=len(contents),
        created_at=now,
    )


@app.delete("/memory/{item_id}")
def delete_memory(item_id: str):
    with db() as c:
        row = c.execute(
            "SELECT path FROM memory_items WHERE id = ?", (item_id,)
        ).fetchone()
        if row:
            try:
                Path(row["path"]).unlink(missing_ok=True)
            except Exception:
                pass
        c.execute("DELETE FROM memory_items WHERE id = ?", (item_id,))
    return {"ok": True}


@app.get("/memory/{item_id}/file")
def download_memory(item_id: str):
    with db() as c:
        row = c.execute(
            "SELECT path, filename FROM memory_items WHERE id = ?", (item_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "Not found.")
    return FileResponse(row["path"], filename=row["filename"])


# ──────────────────────────────────────────────────────────────────────
# Search
# ──────────────────────────────────────────────────────────────────────
def _make_snippet(text: str, query: str, width: int = 180) -> str:
    if not text:
        return ""
    lower = text.lower()
    q = query.lower().strip()
    idx = lower.find(q) if q else -1
    if idx < 0:
        return text[:width].strip() + ("…" if len(text) > width else "")
    start = max(0, idx - width // 2)
    end = min(len(text), idx + len(q) + width // 2)
    snip = text[start:end].strip()
    # bold the match
    if q:
        pattern = re.compile(re.escape(query), re.IGNORECASE)
        snip = pattern.sub(lambda m: f"**{m.group(0)}**", snip)
    prefix = "…" if start > 0 else ""
    suffix = "…" if end < len(text) else ""
    return f"{prefix}{snip}{suffix}"


@app.post("/search", response_model=SearchResponse)
def search(req: SearchRequest):
    q = (req.query or "").strip()
    if not q:
        return SearchResponse(query=q, results=[])

    with db() as c:
        rows = c.execute(
            "SELECT id, filename, kind, size_bytes, text, created_at "
            "FROM memory_items"
        ).fetchall()

    hits: List[SearchHit] = []
    q_lower = q.lower()
    for r in rows:
        name = r["filename"] or ""
        text = r["text"] or ""
        in_name = q_lower in name.lower()
        in_text = q_lower in text.lower()
        if not (in_name or in_text):
            continue
        relevance = 95 if in_name and in_text else (85 if in_name else 70)
        snippet = _make_snippet(text, q) if in_text else f"Match in filename: **{name}**"
        kb = max(1, r["size_bytes"] // 1024)
        meta = f"{r['kind'].upper()} · {kb} KB"
        hits.append(
            SearchHit(
                id=r["id"],
                type=r["kind"],
                title=name,
                path=f"Memory › {r['kind'].upper()} › {_friendly_time(r['created_at'])}",
                snippet=snippet or name,
                meta=meta,
                relevance=relevance,
            )
        )
    hits.sort(key=lambda h: h.relevance, reverse=True)
    return SearchResponse(query=q, results=hits)


# ──────────────────────────────────────────────────────────────────────
# Documents (generate / list / download)
# ──────────────────────────────────────────────────────────────────────
DEFAULT_SECTIONS = [
    "Title Page",
    "Abstract",
    "Introduction",
    "Experimental Setup",
    "Results & Analysis",
    "Conclusion & References",
]


def _doc_body_from_prompt(prompt: str) -> List[tuple[str, str]]:
    """Build a structured (heading, paragraph) list from the prompt.
    Once Ollama is wired in, replace this with real generation per section.
    """
    intro = prompt.strip()
    return [
        ("Abstract",
         f"This document was generated locally on the student's PC. Topic: {intro[:140]}"),
        ("Introduction",
         intro or "Introduction to be generated by the local LLM."),
        ("Experimental Setup",
         "Describe the equipment, materials and procedure used. (Generated locally.)"),
        ("Results & Analysis",
         "Summarise the data collected and interpret the findings. (Generated locally.)"),
        ("Conclusion & References",
         "Summarise the outcomes and list any references used. (Generated locally.)"),
    ]


def _build_docx(title: str, sections: List[tuple[str, str]]) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.size = Pt(20)
    doc.add_paragraph(datetime.now().strftime("Generated %B %d, %Y · My_GPT 4 Students (offline)"))
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(title: str, sections: List[tuple[str, str]]) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, title=title)
    styles = getSampleStyleSheet()
    flow = [
        Paragraph(title, styles["Title"]),
        Paragraph(
            datetime.now().strftime("Generated %B %d, %Y · My_GPT 4 Students (offline)"),
            styles["Italic"],
        ),
        Spacer(1, 18),
    ]
    for heading, body in sections:
        flow.append(Paragraph(heading, styles["Heading2"]))
        flow.append(Paragraph(body.replace("\n", "<br/>"), styles["BodyText"]))
        flow.append(Spacer(1, 10))
    doc.build(flow)
    return buf.getvalue()


@app.get("/documents", response_model=List[DocumentOut])
def list_documents():
    with db() as c:
        rows = c.execute(
            "SELECT id, title, format, size_bytes, created_at FROM documents "
            "ORDER BY created_at DESC"
        ).fetchall()
    return [
        DocumentOut(
            id=r["id"],
            title=r["title"],
            format=r["format"],
            size_bytes=r["size_bytes"],
            created_at=r["created_at"],
            download_url=f"/documents/{r['id']}/download",
            sections=DEFAULT_SECTIONS,
        )
        for r in rows
    ]


@app.post("/documents/generate", response_model=DocumentOut)
def generate_document(req: GenerateDocRequest):
    prompt = (req.prompt or "").strip()
    if not prompt:
        raise HTTPException(400, "Prompt cannot be empty.")
    fmt = (req.format or "docx").lower()
    if fmt not in {"docx", "pdf"}:
        raise HTTPException(400, "Format must be 'docx' or 'pdf'.")

    title = (req.title or prompt.split("\n")[0])[:80].strip() or "Generated Document"
    sections = _doc_body_from_prompt(prompt)

    data = _build_docx(title, sections) if fmt == "docx" else _build_pdf(title, sections)

    doc_id = uuid.uuid4().hex[:12]
    safe_title = re.sub(r"[^A-Za-z0-9]+", "_", title).strip("_") or "document"
    target = EXPORT_DIR / f"{doc_id}__{safe_title}.{fmt}"
    target.write_bytes(data)

    now = datetime.utcnow().isoformat()
    with db() as c:
        c.execute(
            "INSERT INTO documents (id, title, format, path, size_bytes, prompt, created_at) "
            "VALUES (?,?,?,?,?,?,?)",
            (doc_id, title, fmt, str(target), len(data), prompt, now),
        )

    # Drop a record into chat history if a session is provided
    if req.session_id:
        with db() as c:
            if c.execute("SELECT 1 FROM sessions WHERE id = ?", (req.session_id,)).fetchone():
                c.execute(
                    "INSERT INTO messages (session_id, role, content, created_at) VALUES (?,?,?,?)",
                    (req.session_id, "user", f"[Write Doc] {prompt}", now),
                )
                c.execute(
                    "INSERT INTO messages (session_id, role, content, created_at) VALUES (?,?,?,?)",
                    (
                        req.session_id,
                        "assistant",
                        f"Generated **{title}** ({fmt.upper()}). It's saved to your local exports folder.",
                        now,
                    ),
                )
                _bump_session(req.session_id)

    return DocumentOut(
        id=doc_id,
        title=title,
        format=fmt,
        size_bytes=len(data),
        created_at=now,
        download_url=f"/documents/{doc_id}/download",
        sections=[s[0] for s in sections],
    )


@app.get("/documents/{doc_id}/download")
def download_document(doc_id: str):
    with db() as c:
        row = c.execute(
            "SELECT path, title, format FROM documents WHERE id = ?", (doc_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "Not found.")
    fname = re.sub(r"[^A-Za-z0-9]+", "_", row["title"]).strip("_") or "document"
    return FileResponse(row["path"], filename=f"{fname}.{row['format']}")


@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str):
    with db() as c:
        row = c.execute("SELECT path FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if row:
            try:
                Path(row["path"]).unlink(missing_ok=True)
            except Exception:
                pass
        c.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    return {"ok": True}


# ──────────────────────────────────────────────────────────────────────
# Storage info — where data lives, how much is used, how much is free
# ──────────────────────────────────────────────────────────────────────
class StorageInfo(BaseModel):
    data_dir: str
    on_external: bool
    used_bytes: int
    used_human: str
    free_bytes: int
    free_human: str
    total_bytes: int
    total_human: str
    location_file: str
    memory_count: int
    document_count: int


class ModelInfo(BaseModel):
    model_dir: str
    on_external: bool
    location_file: str
    has_models: bool
    model_count: int
    used_bytes: int
    used_human: str
    free_bytes: int
    free_human: str
    total_bytes: int
    total_human: str


@app.get("/storage", response_model=StorageInfo)
def storage_info():
    import shutil
    used = _dir_size(DATA_DIR)
    try:
        usage = shutil.disk_usage(str(DATA_DIR))
        free, total = usage.free, usage.total
    except OSError:
        free, total = 0, 0

    # Heuristic for "external": different drive root than the script's drive.
    on_external = False
    try:
        on_external = Path(DATA_DIR).resolve().drive != ROOT.resolve().drive
        if not on_external:  # *nix fallback: different mount under /Volumes, /media, /mnt
            s = str(DATA_DIR).lower()
            on_external = any(s.startswith(p) for p in ("/volumes/", "/media/", "/mnt/"))
    except Exception:
        pass

    with db() as c:
        mem = c.execute("SELECT COUNT(*) AS n FROM memory_items").fetchone()["n"]
        docs = c.execute("SELECT COUNT(*) AS n FROM documents").fetchone()["n"]

    return StorageInfo(
        data_dir=str(DATA_DIR),
        on_external=on_external,
        used_bytes=used,
        used_human=_human_bytes(used),
        free_bytes=free,
        free_human=_human_bytes(free),
        total_bytes=total,
        total_human=_human_bytes(total),
        location_file=str(LOCATION_FILE),
        memory_count=mem,
        document_count=docs,
    )


@app.get("/model-storage", response_model=ModelInfo)
def model_storage():
    import shutil

    used = _dir_size(MODEL_DIR)
    try:
        usage = shutil.disk_usage(str(MODEL_DIR))
        free, total = usage.free, usage.total
    except OSError:
        free, total = 0, 0

    on_external = False
    try:
        on_external = Path(MODEL_DIR).resolve().drive != ROOT.resolve().drive
        if not on_external:
            s = str(MODEL_DIR).lower()
            on_external = any(s.startswith(p) for p in ("/volumes/", "/media/", "/mnt/"))
    except Exception:
        pass

    model_count = 0
    try:
        model_count = sum(1 for p in MODEL_DIR.rglob("*") if p.is_file())
    except OSError:
        model_count = 0

    return ModelInfo(
        model_dir=str(MODEL_DIR),
        on_external=on_external,
        location_file=str(MODEL_LOCATION_FILE),
        has_models=MODEL_DIR.exists() and any(MODEL_DIR.iterdir()),
        model_count=model_count,
        used_bytes=used,
        used_human=_human_bytes(used),
        free_bytes=free,
        free_human=_human_bytes(free),
        total_bytes=total,
        total_human=_human_bytes(total),
    )


# ──────────────────────────────────────────────────────────────────────
# LLM Configuration (persisted to disk so it survives restarts)
# ──────────────────────────────────────────────────────────────────────
LLM_CONFIG_PATH = DATA_DIR / "llm_config.json"


class LLMConfig(BaseModel):
    ollama_host: str = "http://localhost:11434"
    vision_model: str = "qwen3.5:0.8b"
    reasoning_model: str = "phi4-mini"
    writer_model: str = "llama3.2:1b"


def load_llm_config() -> LLMConfig:
    """Load LLM config from disk, or return defaults."""
    if LLM_CONFIG_PATH.exists():
        try:
            import json
            data = json.loads(LLM_CONFIG_PATH.read_text("utf-8"))
            return LLMConfig(**data)
        except Exception as e:
            print(f"[LLM] Failed to load config: {e}")
    return LLMConfig()


def save_llm_config(cfg: LLMConfig) -> None:
    """Save LLM config to disk."""
    import json
    LLM_CONFIG_PATH.write_text(json.dumps(cfg.model_dump(), ensure_ascii=False), "utf-8")


class LLMTestResult(BaseModel):
    success: bool
    message: str
    available_models: List[str] = []


class LLMStatus(BaseModel):
    ollama_host: str
    vision_model: str
    reasoning_model: str
    writer_model: str
    online: bool
    available_models: List[str]


@app.get("/llm-config", response_model=LLMConfig)
def get_llm_config():
    """Get current LLM configuration."""
    cfg = load_llm_config()
    # Also load from environment if set (env vars override stored config)
    cfg.ollama_host = os.environ.get("OLLAMA_HOST", cfg.ollama_host)
    cfg.vision_model = os.environ.get("OLLAMA_VISION_MODEL", cfg.vision_model)
    cfg.reasoning_model = os.environ.get("OLLAMA_REASONING_MODEL", cfg.reasoning_model)
    cfg.writer_model = os.environ.get("OLLAMA_WRITER_MODEL", cfg.writer_model)
    return cfg


@app.post("/llm-config", response_model=LLMConfig)
def set_llm_config(cfg: LLMConfig):
    """Update LLM configuration and save to disk."""
    save_llm_config(cfg)
    # Update environment vars so they take effect immediately
    os.environ["OLLAMA_HOST"] = cfg.ollama_host
    os.environ["OLLAMA_VISION_MODEL"] = cfg.vision_model
    os.environ["OLLAMA_REASONING_MODEL"] = cfg.reasoning_model
    os.environ["OLLAMA_WRITER_MODEL"] = cfg.writer_model
    return cfg


@app.post("/llm-test", response_model=LLMTestResult)
def test_llm_connection(cfg: LLMConfig):
    """Test connection to Ollama and list available models."""
    try:
        import requests
        
        # Test basic connectivity
        health_url = f"{cfg.ollama_host}/api/tags"
        response = requests.get(health_url, timeout=5)
        
        if response.status_code != 200:
            return LLMTestResult(
                success=False,
                message=f"Ollama not responding (HTTP {response.status_code}). Make sure Ollama is running.",
            )
        
        # Get available models
        data = response.json()
        models = [m.get("name", "") for m in data.get("models", [])]
        
        if not models:
            return LLMTestResult(
                success=True,
                message="Connected to Ollama, but no models found. Pull some models with `ollama pull <model>`.",
                available_models=models,
            )
        
        return LLMTestResult(
            success=True,
            message=f"✓ Connected! Found {len(models)} model(s) available.",
            available_models=sorted(models),
        )
    except requests.ConnectionError:
        return LLMTestResult(
            success=False,
            message=f"Cannot reach Ollama at {cfg.ollama_host}. Is it running?",
        )
    except requests.Timeout:
        return LLMTestResult(
            success=False,
            message=f"Connection to Ollama timed out. Check your host URL and internet.",
        )
    except Exception as e:
        return LLMTestResult(
            success=False,
            message=f"Error testing Ollama: {str(e)}",
        )


@app.get("/llm-status", response_model=LLMStatus)
def get_llm_status():
    """Get current LLM status (config + online check)."""
    cfg = load_llm_config()
    cfg.ollama_host = os.environ.get("OLLAMA_HOST", cfg.ollama_host)
    cfg.vision_model = os.environ.get("OLLAMA_VISION_MODEL", cfg.vision_model)
    cfg.reasoning_model = os.environ.get("OLLAMA_REASONING_MODEL", cfg.reasoning_model)
    cfg.writer_model = os.environ.get("OLLAMA_WRITER_MODEL", cfg.writer_model)
    
    available_models = []
    online = False
    
    try:
        import requests
        response = requests.get(f"{cfg.ollama_host}/api/tags", timeout=2)
        if response.status_code == 200:
            online = True
            data = response.json()
            available_models = sorted([m.get("name", "") for m in data.get("models", [])])
    except Exception:
        pass
    
    return LLMStatus(
        ollama_host=cfg.ollama_host,
        vision_model=cfg.vision_model,
        reasoning_model=cfg.reasoning_model,
        writer_model=cfg.writer_model,
        online=online,
        available_models=available_models,
    )


# ──────────────────────────────────────────────────────────────────────
# Profile (just JSON on disk so it persists across restarts)
# ──────────────────────────────────────────────────────────────────────
PROFILE_PATH = DATA_DIR / "profile.json"


class Profile(BaseModel):
    name: str = "Student"
    career: str = ""
    avatar: Optional[str] = None  # data URL


@app.get("/profile", response_model=Profile)
def get_profile():
    if PROFILE_PATH.exists():
        import json
        try:
            return Profile(**json.loads(PROFILE_PATH.read_text("utf-8")))
        except Exception:
            pass
    return Profile()


@app.put("/profile", response_model=Profile)
def put_profile(p: Profile):
    import json
    PROFILE_PATH.write_text(json.dumps(p.model_dump(), ensure_ascii=False), "utf-8")
    return p


# ──────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    # Honor the env vars set by the Tauri shell when running as a packaged
    # sidecar (MYGPT_HOST / MYGPT_PORT). Falls back to PORT for legacy/dev use.
    host = os.environ.get("MYGPT_HOST", "127.0.0.1")
    port = int(os.environ.get("MYGPT_PORT") or os.environ.get("PORT") or "8000")
    # `reload=True` requires the import-string form and the source on disk —
    # neither is true inside a PyInstaller bundle, so we disable it whenever
    # we're running frozen.
    is_frozen = getattr(sys, "frozen", False)
    if is_frozen:
        uvicorn.run(app, host=host, port=port, log_level="info")
    else:
        uvicorn.run("main:app", host=host, port=port, reload=True)
